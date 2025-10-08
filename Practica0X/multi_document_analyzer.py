"""
Analizador Multi-Documento para Generación de Resúmenes
====================================================

Este script extrae información de múltiples fuentes:
- Archivos PDF (texto e imágenes)
- Documentos DOC/DOCX
- Código fuente (Python, JavaScript, Java, C++, etc.)

Genera un resumen comprehensivo y documentación profesional en formato Markdown.
"""

import os
import ast
import re
import json
from pathlib import Path
from typing import Dict, List, Any, Optional
from datetime import datetime

# Importaciones para diferentes tipos de archivo
try:
    import fitz  # PyMuPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class DocumentAnalyzer:
    """Analizador principal para múltiples tipos de documentos"""
    
    def __init__(self):
        self.resultados = {
            "pdf_content": [],
            "doc_content": [],
            "code_analysis": {},
            "summary": {},
            "timestamp": datetime.now().isoformat()
        }
    
    def extraer_texto_pdf(self, ruta_pdf: str) -> Dict[str, Any]:
        """Extrae texto e imágenes de archivos PDF"""
        if not PDF_AVAILABLE:
            return {"error": "PyMuPDF no está instalado. Instalar con: pip install PyMuPDF"}
        
        print(f"📄 Analizando PDF: {ruta_pdf}")
        
        try:
            doc = fitz.open(ruta_pdf)
            contenido = {
                "archivo": ruta_pdf,
                "paginas": len(doc),
                "texto_completo": "",
                "texto_por_pagina": [],
                "imagenes": []
            }
            
            for num_pagina in range(len(doc)):
                pagina = doc[num_pagina]
                texto_pagina = pagina.get_text()
                
                if texto_pagina.strip():
                    contenido["texto_por_pagina"].append({
                        "pagina": num_pagina + 1,
                        "texto": texto_pagina.strip()
                    })
                    contenido["texto_completo"] += texto_pagina + "\n"
                
                # Contar imágenes (sin extraer por ahora)
                imagenes_pagina = pagina.get_images()
                contenido["imagenes"].append({
                    "pagina": num_pagina + 1,
                    "cantidad_imagenes": len(imagenes_pagina)
                })
            
            doc.close()
            print(f"✅ PDF procesado: {contenido['paginas']} páginas")
            return contenido
            
        except Exception as e:
            return {"error": f"Error procesando PDF: {e}"}
    
    def extraer_texto_docx(self, ruta_doc: str) -> Dict[str, Any]:
        """Extrae texto de documentos DOC/DOCX"""
        if not DOCX_AVAILABLE:
            return {"error": "python-docx no está instalado. Instalar con: pip install python-docx"}
        
        print(f"📝 Analizando documento: {ruta_doc}")
        
        try:
            doc = Document(ruta_doc)
            contenido = {
                "archivo": ruta_doc,
                "parrafos": len(doc.paragraphs),
                "texto_completo": "",
                "parrafos_texto": [],
                "tablas": len(doc.tables),
                "imagenes": len([rel for rel in doc.part.rels.values() 
                               if "image" in rel.target_ref])
            }
            
            # Extraer texto de párrafos
            for i, paragrafo in enumerate(doc.paragraphs):
                if paragrafo.text.strip():
                    contenido["parrafos_texto"].append({
                        "numero": i + 1,
                        "texto": paragrafo.text.strip()
                    })
                    contenido["texto_completo"] += paragrafo.text + "\n"
            
            # Extraer texto de tablas
            tablas_texto = []
            for i, tabla in enumerate(doc.tables):
                tabla_data = []
                for fila in tabla.rows:
                    fila_data = [celda.text.strip() for celda in fila.cells]
                    tabla_data.append(fila_data)
                tablas_texto.append({
                    "tabla": i + 1,
                    "contenido": tabla_data
                })
            
            contenido["tablas_contenido"] = tablas_texto
            
            print(f"✅ Documento procesado: {contenido['parrafos']} párrafos, {contenido['tablas']} tablas")
            return contenido
            
        except Exception as e:
            return {"error": f"Error procesando documento: {e}"}
    
    def analizar_codigo_python(self, ruta_archivo: str) -> Dict[str, Any]:
        """Analiza código fuente Python usando AST"""
        print(f"🐍 Analizando código Python: {ruta_archivo}")
        
        try:
            with open(ruta_archivo, 'r', encoding='utf-8') as archivo:
                codigo = archivo.read()
            
            tree = ast.parse(codigo)
            
            analisis = {
                "archivo": ruta_archivo,
                "tipo": "Python",
                "lineas_codigo": len(codigo.split('\n')),
                "funciones": [],
                "clases": [],
                "imports": [],
                "docstrings": [],
                "comentarios": []
            }
            
            # Extraer información usando AST
            for nodo in ast.walk(tree):
                if isinstance(nodo, ast.FunctionDef):
                    func_info = {
                        "nombre": nodo.name,
                        "linea": nodo.lineno,
                        "argumentos": [arg.arg for arg in nodo.args.args],
                        "docstring": ast.get_docstring(nodo)
                    }
                    analisis["funciones"].append(func_info)
                
                elif isinstance(nodo, ast.ClassDef):
                    class_info = {
                        "nombre": nodo.name,
                        "linea": nodo.lineno,
                        "docstring": ast.get_docstring(nodo)
                    }
                    analisis["clases"].append(class_info)
                
                elif isinstance(nodo, ast.Import):
                    for alias in nodo.names:
                        analisis["imports"].append(alias.name)
                
                elif isinstance(nodo, ast.ImportFrom):
                    modulo = nodo.module or ""
                    for alias in nodo.names:
                        analisis["imports"].append(f"{modulo}.{alias.name}")
            
            # Extraer comentarios usando regex
            comentarios = re.findall(r'#.*', codigo)
            analisis["comentarios"] = [c.strip() for c in comentarios]
            
            print(f"✅ Código Python analizado: {len(analisis['funciones'])} funciones, {len(analisis['clases'])} clases")
            return analisis
            
        except Exception as e:
            return {"error": f"Error analizando código Python: {e}"}
    
    def analizar_codigo_generico(self, ruta_archivo: str) -> Dict[str, Any]:
        """Analiza código de otros lenguajes usando patrones regex"""
        extension = Path(ruta_archivo).suffix.lower()
        
        # Mapeo de extensiones a lenguajes
        lenguajes = {
            '.js': 'JavaScript',
            '.ts': 'TypeScript',
            '.java': 'Java',
            '.cpp': 'C++',
            '.c': 'C',
            '.cs': 'C#',
            '.php': 'PHP',
            '.rb': 'Ruby',
            '.go': 'Go'
        }
        
        lenguaje = lenguajes.get(extension, 'Desconocido')
        print(f"💻 Analizando código {lenguaje}: {ruta_archivo}")
        
        try:
            with open(ruta_archivo, 'r', encoding='utf-8') as archivo:
                codigo = archivo.read()
            
            analisis = {
                "archivo": ruta_archivo,
                "tipo": lenguaje,
                "extension": extension,
                "lineas_codigo": len(codigo.split('\n')),
                "funciones": [],
                "clases": [],
                "imports": [],
                "comentarios": []
            }
            
            # Patrones regex para diferentes elementos
            patrones = {
                'funciones': [
                    r'function\s+(\w+)\s*\(',  # JavaScript
                    r'def\s+(\w+)\s*\(',       # Python
                    r'public\s+\w+\s+(\w+)\s*\(',  # Java
                    r'func\s+(\w+)\s*\(',      # Go
                ],
                'clases': [
                    r'class\s+(\w+)',          # Múltiples lenguajes
                    r'interface\s+(\w+)',      # TypeScript/Java
                ],
                'imports': [
                    r'import\s+([^;]+)',       # JavaScript/TypeScript
                    r'#include\s*[<"]([^>"]+)[>"]',  # C/C++
                    r'using\s+([^;]+)',        # C#
                ],
                'comentarios': [
                    r'//.*',                   # Comentarios de línea
                    r'/\*.*?\*/',             # Comentarios de bloque
                    r'#.*',                   # Python/Shell style
                ]
            }
            
            # Aplicar patrones según el lenguaje
            for categoria, lista_patrones in patrones.items():
                for patron in lista_patrones:
                    matches = re.findall(patron, codigo, re.MULTILINE | re.DOTALL)
                    if matches:
                        analisis[categoria].extend([m.strip() for m in matches])
            
            # Remover duplicados
            for key in ['funciones', 'clases', 'imports', 'comentarios']:
                analisis[key] = list(set(analisis[key]))
            
            print(f"✅ Código {lenguaje} analizado: {len(analisis['funciones'])} funciones, {len(analisis['clases'])} clases")
            return analisis
            
        except Exception as e:
            return {"error": f"Error analizando código {lenguaje}: {e}"}
    
    def analizar_codigo(self, ruta_archivo: str) -> Dict[str, Any]:
        """Punto de entrada para análisis de código"""
        extension = Path(ruta_archivo).suffix.lower()
        
        if extension == '.py':
            return self.analizar_codigo_python(ruta_archivo)
        else:
            return self.analizar_codigo_generico(ruta_archivo)
    
    def generar_resumen_inteligente(self) -> Dict[str, Any]:
        """Genera un resumen inteligente combinando toda la información"""
        print("🧠 Generando resumen inteligente...")
        
        resumen = {
            "proyecto_detectado": "Análisis Multi-Documento",
            "archivos_analizados": {
                "pdf": len(self.resultados["pdf_content"]),
                "doc": len(self.resultados["doc_content"]),
                "codigo": len(self.resultados["code_analysis"])
            },
            "tecnologias_detectadas": set(),
            "funcionalidades_principales": [],
            "estructura_proyecto": {},
            "recomendaciones": []
        }
        
        # Analizar código para detectar tecnologías
        for archivo, analisis in self.resultados["code_analysis"].items():
            if "error" not in analisis:
                resumen["tecnologias_detectadas"].add(analisis["tipo"])
                
                # Detectar frameworks y librerías
                for import_item in analisis.get("imports", []):
                    if any(fw in import_item.lower() for fw in ['flask', 'django', 'fastapi']):
                        resumen["tecnologias_detectadas"].add("Web Framework")
                    if any(fw in import_item.lower() for fw in ['pandas', 'numpy', 'matplotlib']):
                        resumen["tecnologias_detectadas"].add("Data Science")
                    if any(fw in import_item.lower() for fw in ['tensorflow', 'pytorch', 'sklearn']):
                        resumen["tecnologias_detectadas"].add("Machine Learning")
        
        # Convertir set a lista para JSON
        resumen["tecnologias_detectadas"] = list(resumen["tecnologias_detectadas"])
        
        # Generar recomendaciones
        if "Python" in resumen["tecnologias_detectadas"]:
            resumen["recomendaciones"].append("Considerar agregar requirements.txt")
        if "Web Framework" in resumen["tecnologias_detectadas"]:
            resumen["recomendaciones"].append("Documentar endpoints de API")
        if "Data Science" in resumen["tecnologias_detectadas"]:
            resumen["recomendaciones"].append("Incluir ejemplos de datasets")
        
        return resumen
    
    def generar_readme_markdown(self, archivo_salida: str = "README_RESUMEN.md") -> str:
        """Genera un README.md profesional con toda la información"""
        print("📝 Generando README.md profesional...")
        
        resumen = self.generar_resumen_inteligente()
        
        # Generar badges dinámicos basados en tecnologías detectadas
        badges = []
        tech_badges = {
            "Python": "![Python](https://img.shields.io/badge/Python-3776AB?style=for-the-badge&logo=python&logoColor=white)",
            "JavaScript": "![JavaScript](https://img.shields.io/badge/JavaScript-F7DF1E?style=for-the-badge&logo=javascript&logoColor=black)",
            "Java": "![Java](https://img.shields.io/badge/Java-ED8B00?style=for-the-badge&logo=java&logoColor=white)",
            "C++": "![C++](https://img.shields.io/badge/C++-00599C?style=for-the-badge&logo=c%2B%2B&logoColor=white)",
            "Web Framework": "![Web](https://img.shields.io/badge/Web-Framework-green?style=for-the-badge)",
            "Data Science": "![Data Science](https://img.shields.io/badge/Data-Science-orange?style=for-the-badge)",
            "Machine Learning": "![ML](https://img.shields.io/badge/Machine-Learning-red?style=for-the-badge)"
        }
        
        for tech in resumen["tecnologias_detectadas"]:
            if tech in tech_badges:
                badges.append(tech_badges[tech])
        
        markdown_content = f"""<div align="center">

# 📊 **Análisis Multi-Documento del Proyecto**

## **Resumen Automático Generado**

**Análisis comprehensivo de documentación (PDF/DOC) y código fuente para generar documentación técnica completa.**

{' '.join(badges)}

![Análisis](https://img.shields.io/badge/Análisis-Completo-brightgreen?style=for-the-badge)
![Archivos](https://img.shields.io/badge/Archivos-{resumen['archivos_analizados']['pdf'] + resumen['archivos_analizados']['doc'] + resumen['archivos_analizados']['codigo']}-blue?style=for-the-badge)

</div>

---

## 📋 **Contenido**
- [🎯 Resumen Ejecutivo](#-resumen-ejecutivo)
- [📁 Archivos Analizados](#-archivos-analizados)
- [💻 Análisis de Código](#-análisis-de-código)
- [📄 Documentación Encontrada](#-documentación-encontrada)
- [🛠️ Tecnologías Detectadas](#️-tecnologías-detectadas)
- [💡 Recomendaciones](#-recomendaciones)

## 🎯 **Resumen Ejecutivo**

Este proyecto ha sido analizado automáticamente procesando **{resumen['archivos_analizados']['pdf'] + resumen['archivos_analizados']['doc'] + resumen['archivos_analizados']['codigo']} archivos** de diferentes tipos:

- **📄 Documentos PDF:** {resumen['archivos_analizados']['pdf']} archivos
- **📝 Documentos DOC/DOCX:** {resumen['archivos_analizados']['doc']} archivos  
- **💻 Archivos de código:** {resumen['archivos_analizados']['codigo']} archivos

### **🔍 Tecnologías Principales Detectadas:**
"""
        
        for tech in resumen["tecnologias_detectadas"]:
            markdown_content += f"- **{tech}**\n"
        
        # Sección de análisis de código
        markdown_content += "\n## 💻 **Análisis de Código**\n\n"
        
        for archivo, analisis in self.resultados["code_analysis"].items():
            if "error" not in analisis:
                markdown_content += f"### 📂 `{os.path.basename(archivo)}`\n\n"
                markdown_content += f"- **Lenguaje:** {analisis['tipo']}\n"
                markdown_content += f"- **Líneas de código:** {analisis['lineas_codigo']}\n"
                markdown_content += f"- **Funciones:** {len(analisis.get('funciones', []))}\n"
                markdown_content += f"- **Clases:** {len(analisis.get('clases', []))}\n"
                markdown_content += f"- **Imports:** {len(analisis.get('imports', []))}\n\n"
                
                # Listar funciones principales
                if analisis.get('funciones'):
                    markdown_content += "**Funciones principales:**\n"
                    for func in analisis['funciones'][:5]:  # Mostrar solo las primeras 5
                        if isinstance(func, dict):
                            markdown_content += f"- `{func.get('nombre', 'Unknown')}()` (línea {func.get('linea', '?')})\n"
                        else:
                            markdown_content += f"- `{func}()`\n"
                    markdown_content += "\n"
        
        # Sección de documentación
        markdown_content += "## 📄 **Documentación Encontrada**\n\n"
        
        # Información de PDFs
        for pdf_info in self.resultados["pdf_content"]:
            if "error" not in pdf_info:
                markdown_content += f"### 📄 {os.path.basename(pdf_info['archivo'])}\n\n"
                markdown_content += f"- **Páginas:** {pdf_info['paginas']}\n"
                markdown_content += f"- **Imágenes:** {sum(img['cantidad_imagenes'] for img in pdf_info['imagenes'])}\n"
                markdown_content += f"- **Contenido:** Documentación técnica y guías\n\n"
        
        # Información de documentos DOC
        for doc_info in self.resultados["doc_content"]:
            if "error" not in doc_info:
                markdown_content += f"### 📝 {os.path.basename(doc_info['archivo'])}\n\n"
                markdown_content += f"- **Párrafos:** {doc_info['parrafos']}\n"
                markdown_content += f"- **Tablas:** {doc_info['tablas']}\n"
                markdown_content += f"- **Imágenes:** {doc_info['imagenes']}\n\n"
        
        # Recomendaciones
        markdown_content += "## 💡 **Recomendaciones**\n\n"
        for recomendacion in resumen["recomendaciones"]:
            markdown_content += f"- ✅ {recomendacion}\n"
        
        if not resumen["recomendaciones"]:
            markdown_content += "- ✅ El proyecto parece estar bien estructurado\n"
            markdown_content += "- ✅ Considerar agregar más documentación inline\n"
            markdown_content += "- ✅ Revisar la cobertura de tests\n"
        
        # Footer
        markdown_content += f"""
## 📊 **Estadísticas del Análisis**

| Métrica | Valor |
|---------|-------|
| Archivos PDF procesados | {resumen['archivos_analizados']['pdf']} |
| Documentos DOC procesados | {resumen['archivos_analizados']['doc']} |
| Archivos de código analizados | {resumen['archivos_analizados']['codigo']} |
| Tecnologías detectadas | {len(resumen['tecnologias_detectadas'])} |
| Fecha de análisis | {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} |

<div align="center">

---

### 🎯 **Análisis generado automáticamente** 🤖

**Creado con ❤️ usando Python y análisis inteligente de documentos**

</div>
"""
        
        # Guardar archivo
        with open(archivo_salida, 'w', encoding='utf-8') as archivo:
            archivo.write(markdown_content)
        
        print(f"✅ README generado: {archivo_salida}")
        return archivo_salida
    
    def procesar_multiples_archivos(self, archivos: List[str]) -> str:
        """Procesa múltiples archivos y genera resumen completo"""
        print("🚀 Iniciando análisis multi-documento...")
        
        for archivo in archivos:
            if not os.path.exists(archivo):
                print(f"⚠️  Archivo no encontrado: {archivo}")
                continue
            
            extension = Path(archivo).suffix.lower()
            
            if extension == '.pdf':
                resultado = self.extraer_texto_pdf(archivo)
                self.resultados["pdf_content"].append(resultado)
            
            elif extension in ['.docx', '.doc']:
                resultado = self.extraer_texto_docx(archivo)
                self.resultados["doc_content"].append(resultado)
            
            elif extension in ['.py', '.js', '.ts', '.java', '.cpp', '.c', '.cs', '.php', '.rb', '.go']:
                resultado = self.analizar_codigo(archivo)
                self.resultados["code_analysis"][archivo] = resultado
            
            else:
                print(f"⚠️  Tipo de archivo no soportado: {archivo}")
        
        # Generar resumen final
        readme_file = self.generar_readme_markdown()
        
        # Guardar resultados completos en JSON
        with open("analisis_completo.json", 'w', encoding='utf-8') as f:
            json.dump(self.resultados, f, indent=2, ensure_ascii=False)
        
        print("\n" + "="*70)
        print("🎉 ANÁLISIS MULTI-DOCUMENTO COMPLETADO")
        print("="*70)
        print(f"📄 README generado: {readme_file}")
        print(f"📊 Datos completos: analisis_completo.json")
        print(f"📁 Archivos procesados: {len(archivos)}")
        print("="*70)
        
        return readme_file

def main():
    """Función principal de demostración"""
    print("🔍 Analizador Multi-Documento")
    print("="*50)
    
    # Ejemplo de uso
    archivos_ejemplo = [
        "documento.pdf",
        "manual.docx", 
        "codigo_principal.py",
        "script.js"
    ]
    
    print("Archivos de ejemplo que puedes analizar:")
    for archivo in archivos_ejemplo:
        print(f"  - {archivo}")
    
    print("\nPara usar el analizador:")
    print("1. Instalar dependencias: pip install PyMuPDF python-docx")
    print("2. Crear instancia: analyzer = DocumentAnalyzer()")
    print("3. Procesar archivos: analyzer.procesar_multiples_archivos([lista_archivos])")

if __name__ == "__main__":
    main()